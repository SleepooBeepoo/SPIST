"""Document Processor Module

This module handles the extraction of questions from document files (.docx, .pdf)
with improved error handling and validation.
"""
import os
import re
import json
import docx
import PyPDF2
import tempfile
from werkzeug.utils import secure_filename
from typing import List, Dict, Tuple, Optional, Any

# Check which OpenAI library version is being used
try:
    import openai
    from openai import OpenAI
    USING_NEW_OPENAI = True
except ImportError:
    import openai
    USING_NEW_OPENAI = False

class DocumentProcessor:
    """Class to handle processing document files and extracting questions"""
    
    def __init__(self):
        """Initialize the document processor"""
        # Get OpenAI API key from environment variable
        self.api_key = os.environ.get('OPENAI_API_KEY')
        
        # Initialize OpenAI client if API key is available
        if self.api_key:
            if USING_NEW_OPENAI:
                self.client = OpenAI(api_key=self.api_key)
            else:
                openai.api_key = self.api_key
    
    def save_uploaded_file(self, file, upload_dir: str) -> str:
        """Save the uploaded file to the specified directory
        
        Args:
            file: The uploaded file object
            upload_dir: The directory to save the file to
            
        Returns:
            str: The path to the saved file
        """
        # Create the upload directory if it doesn't exist
        if not os.path.exists(upload_dir):
            os.makedirs(upload_dir, exist_ok=True)
        
        # Secure the filename to prevent directory traversal attacks
        filename = secure_filename(file.filename)
        file_path = os.path.join(upload_dir, filename)
        
        # Save the file
        file.save(file_path)
        
        return file_path
    
    def process_file(self, file_path: str, use_ai: bool = False) -> Tuple[List[Dict[str, Any]], Optional[str]]:
        """Process the file and extract questions
        
        Args:
            file_path: The path to the file to process
            use_ai: Whether to use AI to extract questions
            
        Returns:
            Tuple[List[Dict], Optional[str]]: A tuple containing the list of extracted questions
                and an optional error message
        """
        try:
            # Print debug information
            print(f"Processing file: {file_path}")
            print(f"Use AI: {use_ai}")
            
            # Check if the file exists
            if not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                return [], f"File not found: {file_path}"
            
            # Get the file extension
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            print(f"File extension: {ext}")
            
            # Extract text based on file type
            if ext == '.docx':
                print("Processing DOCX file")
                text = self._extract_text_from_docx(file_path)
            elif ext == '.pdf':
                print("Processing PDF file")
                text = self._extract_text_from_pdf(file_path)
            else:
                print(f"Unsupported file type: {ext}")
                return [], f"Unsupported file type: {ext}"
            
            # Check if text was extracted successfully
            if not text:
                print("Failed to extract text from the document")
                return [], "Failed to extract text from the document"
            
            print(f"Successfully extracted {len(text)} characters of text")
            
            # Use AI to extract questions if enabled and API key is available
            if use_ai and self.api_key:
                try:
                    print("Using AI to extract questions")
                    questions, error = self._extract_questions_with_ai(text)
                    if error:
                        print(f"AI extraction error: {error}. Falling back to regex parsing.")
                        # Continue with regex parsing below
                    else:
                        print(f"AI successfully extracted {len(questions)} questions")
                        return questions, None
                except Exception as e:
                    print(f"AI extraction failed: {str(e)}. Falling back to regex parsing.")
                    # Fall back to regex parsing if AI fails
            else:
                print("Using regex parsing to extract questions")
        except Exception as e:
            print(f"Unexpected error in process_file: {str(e)}")
            return [], f"Unexpected error processing file: {str(e)}"
        
        # Parse the text to extract questions using regex
        questions = self._parse_text_for_questions(text)
        
        return questions, None
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a .docx file
        
        Args:
            file_path: The path to the .docx file
            
        Returns:
            str: The extracted text
        """
        try:
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(' | '.join(row_text))
            
            # Combine all text
            all_text = '\n'.join(paragraphs + table_text)
            
            return all_text
        except Exception as e:
            print(f"Error extracting text from DOCX: {str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a .pdf file
        
        Args:
            file_path: The path to the .pdf file
            
        Returns:
            str: The extracted text
        """
        try:
            # Open the PDF file
            with open(file_path, 'rb') as file:
                # Create a PDF reader object
                reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                text = []
                for page in reader.pages:
                    text.append(page.extract_text())
                
                # Combine all text
                all_text = '\n'.join(text)
                
                return all_text
        except Exception as e:
            print(f"Error extracting text from PDF: {str(e)}")
            return ""
    
    def _parse_text_for_questions(self, text: str) -> List[Dict[str, Any]]:
        """Parse the extracted text to identify questions
        
        Args:
            text: The text to parse
            
        Returns:
            List[Dict]: A list of dictionaries containing question data
        """
        questions = []
        
        # Split the text into lines
        lines = text.split('\n')
        
        # Initialize variables
        current_question = None
        current_options = []
        question_number = 0
        
        # Regular expressions for identifying questions and options
        question_pattern = re.compile(r'^\s*(?:\d+\.?|[Qq]uestion\s*\d*:?|[Qq]:?)\s*(.+)$')
        option_pattern = re.compile(r'^\s*([A-Za-z])\s*[.)]\s*(.+)$')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a question
            question_match = question_pattern.match(line)
            if question_match:
                # If we were processing a previous question, add it to the list
                if current_question:
                    self._add_question_to_list(questions, current_question, current_options)
                
                # Start a new question
                question_number += 1
                current_question = question_match.group(1).strip()
                current_options = []
                continue
            
            # Check if this line is an option
            option_match = option_pattern.match(line)
            if option_match and current_question:
                option_text = option_match.group(2).strip()
                current_options.append(option_text)
                continue
            
            # If we're in a question but this line isn't an option, it might be part of the question
            if current_question:
                # Check if this line indicates the correct answer
                if line.lower().startswith(('answer:', 'correct:', 'correct answer:')):
                    # Process the correct answer
                    answer_text = line.split(':', 1)[1].strip()
                    self._process_correct_answer(questions, current_question, current_options, answer_text)
                    
                    # Reset for the next question
                    current_question = None
                    current_options = []
                    continue
                    
        # Add the last question if there is one
        if current_question:
            self._add_question_to_list(questions, current_question, current_options)
            
        return questions
        
    def _extract_questions_with_ai(self, text: str) -> Tuple[List[Dict[str, Any]], Optional[str]]:
        """Use AI to extract questions from text content
        
        Args:
            text: The text content to extract questions from
            
        Returns:
            Tuple[List[Dict], Optional[str]]: A tuple containing the list of extracted questions
                and an optional error message
        """
        try:
            # Prepare the prompt for the AI
            prompt = f"""Extract questions from the following document text. For each question, determine:
1. The question text
2. The question type (multiple_choice, true_false, identification, or essay)
3. The options (for multiple choice questions)
4. The correct answer
5. The point value (default to 1.0)

Format your response as a JSON array of question objects.

Here's the document text:

{text}
"""
            
            # Call the OpenAI API
            if USING_NEW_OPENAI:
                response = self.client.chat.completions.create(
                    model="gpt-3.5-turbo-16k",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant that extracts quiz questions from documents."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )
                result = response.choices[0].message.content
            else:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo-16k",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant that extracts quiz questions from documents."},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.3,
                    max_tokens=4000
                )
                result = response.choices[0].message['content']
            
            # Extract the JSON part from the response
            json_match = re.search(r'\[\s*\{.+\}\s*\]', result, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
            else:
                json_str = result
            
            # Parse the JSON
            try:
                questions = json.loads(json_str)
                
                # Validate and clean up the questions
                cleaned_questions = []
                for q in questions:
                    # Ensure required fields
                    if 'question_text' not in q or not q['question_text'].strip():
                        continue
                    
                    # Set default question type if missing
                    if 'question_type' not in q or q['question_type'] not in ['multiple_choice', 'true_false', 'identification', 'essay']:
                        q['question_type'] = 'essay'
                    
                    # Set default points if missing
                    if 'points' not in q or not isinstance(q['points'], (int, float)):
                        q['points'] = 1.0
                    
                    # Handle type-specific fields
                    if q['question_type'] == 'multiple_choice':
                        if 'options' not in q or not isinstance(q['options'], list) or len(q['options']) < 2:
                            q['options'] = ['Option 1', 'Option 2', 'Option 3', 'Option 4']
                        
                        if 'correct_answer' not in q or not isinstance(q['correct_answer'], (str, int)):
                            q['correct_answer'] = '0'
                        else:
                            # Ensure correct_answer is a string
                            q['correct_answer'] = str(q['correct_answer'])
                    
                    elif q['question_type'] == 'true_false':
                        if 'correct_answer' not in q or q['correct_answer'] not in ['true', 'false']:
                            q['correct_answer'] = 'true'
                    
                    elif q['question_type'] == 'identification':
                        if 'correct_answer' not in q:
                            q['correct_answer'] = ''
                    
                    elif q['question_type'] == 'essay':
                        if 'word_limit' not in q or not isinstance(q['word_limit'], int):
                            q['word_limit'] = 500
                        
                        if 'correct_answer' not in q:
                            q['correct_answer'] = ''
                    
                    cleaned_questions.append(q)
                
                return cleaned_questions, None
            except json.JSONDecodeError as e:
                print(f"JSON parsing error: {str(e)}")
                print(f"Attempted to parse: {json_str}")
                # Fall back to regex extraction with error message
                error_msg = f"JSON parsing error: {str(e)}"
                questions = self._parse_text_for_questions(text)
                return questions, error_msg
        
        except Exception as e:
            print(f"AI extraction error: {str(e)}")
            # Fall back to regex extraction with error message
            error_msg = f"AI extraction error: {str(e)}"
            questions = self._parse_text_for_questions(text)
            return questions, error_msg
    
    def _add_question_to_list(self, questions: List[Dict[str, Any]], question_text: str, options: List[str]) -> None:
        """Add a question to the list of questions
        
        Args:
            questions: The list to add the question to
            question_text: The text of the question
            options: The list of options for the question
        """
        # Determine question type based on options
        if len(options) >= 2:
            # Check if it's a true/false question
            if len(options) == 2 and (
                (options[0].lower() in ['true', 't'] and options[1].lower() in ['false', 'f']) or
                (options[0].lower() in ['false', 'f'] and options[1].lower() in ['true', 't'])
            ):
                question_type = 'true_false'
                correct_answer = '0' if options[0].lower() in ['true', 't'] else '1'
            else:
                question_type = 'multiple_choice'
                # Default to first option as correct (will be updated later if specified)
                correct_answer = '0'
        elif question_text.lower().endswith('true or false?') or question_text.lower().endswith('true or false.'):
            # True/False without options
            question_type = 'true_false'
            correct_answer = 'true'  # Default to true
        elif '___' in question_text or '...' in question_text or '_____' in question_text:
            # Identification/Fill-in-the-blank
            question_type = 'identification'
            correct_answer = ''
        else:
            # Default to identification if no options
            question_type = 'identification'
            correct_answer = ''
        
        # Create the question dictionary
        question = {
            'question_text': question_text,
            'question_type': question_type,
            'options': options if options and question_type == 'multiple_choice' else None,
            'correct_answer': correct_answer,
            'points': 1.0  # Default points
        }
        
        # Add word limit for essay questions
        if question_type == 'essay':
            question['word_limit'] = 500
        
        questions.append(question)
    
    def _process_correct_answer(self, questions: List[Dict[str, Any]], question_text: str, 
                               options: List[str], answer_text: str) -> None:
        """Process the correct answer for a question
        
        Args:
            questions: The list of questions
            question_text: The text of the question
            options: The list of options for the question
            answer_text: The text indicating the correct answer
        """
        # Create the question dictionary
        question = {
            'question_text': question_text,
            'options': options if options else None,
            'points': 1.0  # Default points
        }
        
        # Determine question type and correct answer based on options and answer text
        # Check if it's a true/false question based on the question text
        if question_text.lower().endswith('true or false') or question_text.lower().endswith('true/false') or \
           'true or false:' in question_text.lower() or 'true/false:' in question_text.lower() or \
           question_text.lower().strip().startswith('true or false:') or \
           question_text.lower().strip().startswith('this') or question_text.lower().strip().startswith('that') or \
           question_text.lower().strip().startswith('the following') or question_text.lower().strip().startswith('the statement'):
            # True/False question
            question['question_type'] = 'true_false'
            
            # Normalize the answer
            if answer_text.lower() in ['true', 't', '1', 'yes', 'y', 'this', 'that', 'correct', 'right']:
                question['correct_answer'] = 'true'
            else:
                question['correct_answer'] = 'false'
                
        elif options and len(options) >= 2:
            if len(options) == 2 and (
                (options[0].lower() in ['true', 't'] and options[1].lower() in ['false', 'f']) or
                (options[0].lower() in ['false', 'f'] and options[1].lower() in ['true', 't'])
            ):
                # True/False question
                question['question_type'] = 'true_false'
                
                # Determine correct answer
                if answer_text.lower() in ['true', 't', 'a', 'this', 'that', 'correct', 'right']:
                    question['correct_answer'] = '0' if options[0].lower() in ['true', 't'] else '1'
                else:
                    question['correct_answer'] = '1' if options[0].lower() in ['true', 't'] else '0'
            else:
                # Multiple choice question
                question['question_type'] = 'multiple_choice'
                
                # Try to determine the correct answer from the answer text
                answer_text = answer_text.strip().upper()
                if answer_text in 'ABCDEFGHIJKLMNOPQRSTUVWXYZ':
                    # Convert letter to index (A=0, B=1, etc.)
                    index = ord(answer_text) - ord('A')
                    if 0 <= index < len(options):
                        question['correct_answer'] = str(index)
                    else:
                        question['correct_answer'] = '0'  # Default to first option
                elif answer_text.isdigit():
                    # Numeric answer
                    index = int(answer_text) - 1
                    if 0 <= index < len(options):
                        question['correct_answer'] = str(index)
                    else:
                        question['correct_answer'] = '0'  # Default to first option
                else:
                    # Try to match the answer text with an option
                    found_match = False
                    for i, option in enumerate(options):
                        if answer_text.lower() in option.lower():
                            question['correct_answer'] = str(i)
                            found_match = True
                            break
                    
                    if not found_match:
                        question['correct_answer'] = '0'  # Default to first option
        elif answer_text.lower() in ['true', 'false']:
            # True/False without options
            question['question_type'] = 'true_false'
            question['correct_answer'] = answer_text.lower()
        else:
            # Identification or essay question
            # First check for explicit labels in the question text
            if 'essay' in question_text.lower() or 'explain' in question_text.lower() or 'describe' in question_text.lower() or 'discuss' in question_text.lower():
                # Explicitly labeled as essay
                question['question_type'] = 'essay'
                question['word_limit'] = 500
            elif 'identification' in question_text.lower() or 'identify' in question_text.lower() or 'name' in question_text.lower() or 'what is' in question_text.lower():
                # Explicitly labeled as identification
                question['question_type'] = 'identification'
            # If no explicit label, use heuristics
            elif len(question_text) > 200 or '?' not in question_text:
                # Likely an essay based on length
                question['question_type'] = 'essay'
                question['word_limit'] = 500
            else:
                # Default to identification
                question['question_type'] = 'identification'
            
            question['correct_answer'] = answer_text
        
        questions.append(question)